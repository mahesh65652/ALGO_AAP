import io
import sqlite3
import hashlib
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit.components.v1 as components

# ── PAGE CONFIG ────────────────────────────────────────────────────
st.set_page_config(
    page_title="ઓમ જનસેવા & ઓનલાઈન સોલ્યુશન સેન્ટર",
    page_icon="💻",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ── 3CX STYLE ULTRA MODERN DARK NEON CSS ──────────────────────────
st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(180deg, #020b18 0%, #08142b 100%) !important;
        color: #ffffff !important;
        max-width: 100vw;
        overflow-x: hidden !important;
    }
    
    h1, h2, h3, h4, h5, h6, span, label, p, .stMarkdown {
        color: #ffffff !important;
    }
    
    .main-title {
        color: #40a9ff;
        font-size: 30px;
        font-weight: 800;
        text-align: center;
        margin-bottom: 5px;
    }
    .sub-title {
        color: #8c8c8c;
        font-size: 15px;
        text-align: center;
        margin-bottom: 25px;
    }

    .card-box {
        background: rgba(13, 27, 54, 0.7);
        border: 1px solid #173b70;
        border-radius: 12px;
        padding: 18px;
        text-align: center;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    .card-box:hover {
        border-color: #1890ff;
        box-shadow: 0 0 15px rgba(24, 144, 255, 0.4);
    }
    .card-icon { font-size: 26px; margin-bottom: 8px; }
    .card-title { color: #1890ff !important; font-size: 17px; font-weight: bold; margin-bottom: 6px; }
    .card-desc { color: #a6c5e8 !important; font-size: 13px; line-height: 1.4; }

    section[data-testid="stSidebar"] {
        background-color: #051024 !important;
        border-right: 1px solid #173b70;
    }

    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        color: #ffffff !important;
        background-color: #0b1930 !important;
        border: 1px solid #173b70 !important;
        border-radius: 8px !important;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px 8px 0px 0px;
        padding: 10px 20px;
        background-color: #091a38;
        color: #ffffff !important;
        border: 1px solid #173b70;
    }
    .stTabs [aria-selected="true"] {
        background-color: #1890ff !important;
        color: #ffffff !important;
    }

    .btn-3d-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
        gap: 12px;
        margin: 15px 0;
    }
    .btn-3d {
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 12px 16px;
        font-size: 14px;
        font-weight: 700;
        color: #ffffff !important;
        text-decoration: none !important;
        border-radius: 10px;
        box-shadow: 0 4px 0 rgba(0, 0, 0, 0.4);
        transition: all 0.15s ease-in-out;
    }
    .btn-3d:active { transform: translateY(2px); box-shadow: 0 2px 0 rgba(0,0,0,0.4); }
    .btn-blue { background: linear-gradient(145deg, #1e88e5, #1565c0); }
    .btn-green { background: linear-gradient(145deg, #43a047, #2e7d32); }
    .btn-purple { background: linear-gradient(145deg, #8e24aa, #6a1b9a); }
    .btn-orange { background: linear-gradient(145deg, #fb8c00, #ef6c00); }
    .btn-red { background: linear-gradient(145deg, #e53935, #c62828); }

    .email-result-card {
        background: #091a38;
        border: 1px solid #1890ff;
        border-radius: 10px;
        padding: 15px;
        margin-top: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# ── GUJARAT DISTRICT & TALUKA GOVT EMAIL DATABASE ─────────────────
GOVT_DISTRICT_EMAILS = {
    "મોરબી (Morbi)": {
        "કલેક્ટર કચેરી": "collector-mor@gujarat.gov.in",
        "SDM કચેરી (મોરબી)": "sdm-morbi-raj@gujarat.gov.in",
        "મામલતદાર (મોરબી તાલુકા)": "mam-morbi@gujarat.gov.in",
        "મામલતદાર (ટંકારા તાલુકા)": "mam-tankara@gujarat.gov.in",
        "મામલતદાર (વાંકાનેર તાલુકા)": "mam-wankaner@gujarat.gov.in",
        "મામલતદાર (હળવદ તાલુકા)": "mam-halvad@gujarat.gov.in",
        "મામલતદાર (માળીયા મિયાણા)": "mam-maliya@gujarat.gov.in"
    },
    "રાજકોટ (Rajkot)": {
        "કલેક્ટર કચેરી": "collector-raj@gujarat.gov.in",
        "SDM કચેરી (રાજકોટ શહેર)": "sdm-rajkot@gujarat.gov.in",
        "મામલતદાર (રાજકોટ શહેર)": "mam-rajkot@gujarat.gov.in",
        "મામલતદાર (ગોંડલ તાલુકા)": "mam-gondal@gujarat.gov.in",
        "મામલતદાર (જેતપુર તાલુકા)": "mam-jetpur@gujarat.gov.in",
        "મામલતદાર (જસદણ તાલુકા)": "mam-jasdan@gujarat.gov.in",
        "મામલતદાર (ધોરાજી તાલુકા)": "mam-dhoraji@gujarat.gov.in"
    },
    "સુરત (Surat)": {
        "કલેક્ટર કચેરી": "collector-sur@gujarat.gov.in",
        "SMC કમિશનર": "commissioner@suratmunicipal.gov.in",
        "મામલતદાર (ચોર્યાસી તાલુકા)": "mam-choryasi@gujarat.gov.in",
        "મામલતદાર (કામરેજ તાલુકા)": "mam-kamrej@gujarat.gov.in",
        "મામલતદાર (ઓલપાડ તાલુકા)": "mam-olpad@gujarat.gov.in",
        "મામલતદાર (બારડોલી તાલુકા)": "mam-bardoli@gujarat.gov.in"
    },
    "અમદાવાદ (Ahmedabad)": {
        "કલેક્ટર કચેરી": "collector-ahd@gujarat.gov.in",
        "AMC કમિશનર": "mc@ahmedabadcity.gov.in",
        "મામલતદાર (દસ્ક્રોઈ તાલુકા)": "mam-daskroi@gujarat.gov.in",
        "મામલતદાર (સાણંદ તાલુકા)": "mam-sanand@gujarat.gov.in",
        "મામલતદાર (ધોળકા તાલુકા)": "mam-dholka@gujarat.gov.in"
    },
    "વડોદરા (Vadodara)": {
        "કલેક્ટર કચેરી": "collector-vad@gujarat.gov.in",
        "SDM કચેરી": "sdm-vadodara@gujarat.gov.in",
        "મામલતદાર (ડભોઈ તાલુકા)": "mam-dabhoi@gujarat.gov.in",
        "મામલતદાર (વાઘોડિયા તાલુકા)": "mam-waghodia@gujarat.gov.in"
    },
    "જામનગર (Jamnagar)": {
        "કલેક્ટર કચેરી": "collector-jam@gujarat.gov.in",
        "મામલતદાર (જામનગર તાલુકા)": "mam-jamnagar@gujarat.gov.in",
        "મામલતદાર (લાલપુર તાલુકા)": "mam-lalpur@gujarat.gov.in",
        "મામલતદાર (ધ્રોલ તાલુકા)": "mam-dhrol@gujarat.gov.in"
    },
    "જૂનાગઢ (Junagadh)": {
        "કલેક્ટર કચેરી": "collector-jun@gujarat.gov.in",
        "મામલતદાર (કેશોદ તાલુકા)": "mam-keshod@gujarat.gov.in",
        "મામલતદાર (માણાવદર તાલુકા)": "mam-manavadar@gujarat.gov.in"
    },
    "ભાવનગર (Bhavnagar)": {
        "કલેક્ટર કચેરી": "collector-bha@gujarat.gov.in",
        "મામલતદાર (મહુવા તાલુકા)": "mam-mahuva@gujarat.gov.in",
        "મામલતદાર (પાલિતાણા તાલુકા)": "mam-palitana@gujarat.gov.in"
    },
    "કચ્છ (Kutch)": {
        "કલેક્ટર કચેરી": "collector-kuch@gujarat.gov.in",
        "મામલતદાર (ગાંધીધામ તાલુકા)": "mam-gandhidham@gujarat.gov.in",
        "મામલતદાર (અંજાર તાલુકા)": "mam-anjar@gujarat.gov.in",
        "મામલતદાર (માંડવી તાલુકા)": "mam-mandvi@gujarat.gov.in"
    }
}

# ── DATABASE SETUP ─────────────────────────────────────────────────
DB_FILE = "om_janseva_drafts.db"

def init_db():
    with sqlite3.connect(DB_FILE) as conn:
        c = conn.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS users (
                username TEXT PRIMARY KEY,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS drafts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL,
                doc_type TEXT NOT NULL,
                service_category TEXT,
                office_name TEXT,
                applicant_name TEXT,
                contact_no TEXT,
                reference_no TEXT,
                doc_body TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        c.execute("SELECT COUNT(*) FROM users")
        if c.fetchone()[0] == 0:
            admin_pass = hashlib.sha256("admin123".encode()).hexdigest()
            c.execute("INSERT INTO users VALUES (?, ?, ?)", ("admin", admin_pass, "admin"))

init_db()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_user(username, password):
    with sqlite3.connect(DB_FILE) as conn:
        c = conn.cursor()
        c.execute("SELECT password_hash, role FROM users WHERE username = ?", (username,))
        res = c.fetchone()
        if res and res[0] == hash_password(password):
            return res[1]
    return None

def register_user(username, password, role="client"):
    try:
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            c.execute("INSERT INTO users VALUES (?, ?, ?)", (username, hash_password(password), role))
            return True
    except sqlite3.IntegrityError:
        return False

def save_user_draft(username, doc_type, service_category, office_name, applicant_name, contact_no, reference_no, doc_body):
    with sqlite3.connect(DB_FILE) as conn:
        c = conn.cursor()
        c.execute("""
            INSERT INTO drafts (username, doc_type, service_category, office_name, applicant_name, contact_no, reference_no, doc_body)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (username, doc_type, service_category, office_name, applicant_name, contact_no, reference_no, doc_body))

def get_user_drafts(username, search_query="", is_admin=False):
    with sqlite3.connect(DB_FILE) as conn:
        c = conn.cursor()
        query = "%" + search_query + "%"
        if is_admin:
            c.execute("SELECT * FROM drafts WHERE applicant_name LIKE ? OR doc_type LIKE ? ORDER BY created_at DESC", (query, query))
        else:
            c.execute("SELECT * FROM drafts WHERE username = ? AND (applicant_name LIKE ? OR doc_type LIKE ?) ORDER BY created_at DESC", (username, query, query))
        return c.fetchall()

def delete_user_draft(draft_id, username, is_admin=False):
    with sqlite3.connect(DB_FILE) as conn:
        c = conn.cursor()
        if is_admin:
            c.execute("DELETE FROM drafts WHERE id = ?", (draft_id,))
        else:
            c.execute("DELETE FROM drafts WHERE id = ? AND username = ?", (draft_id, username))

# ── SESSION STATE ──────────────────────────────────────────────────
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False
if "username" not in st.session_state:
    st.session_state["username"] = ""
if "role" not in st.session_state:
    st.session_state["role"] = ""

# ── SIDEBAR (ONLY LOGIN & LINKS) ────────────────────────────────────
with st.sidebar:
    st.title("💻 ઓમ જનસેવા સેન્ટર")
    st.caption("સંચાલક: મહેશભાઈ રામાવત")

    with st.expander("🔐 લોગિન / એકાઉન્ટ", expanded=not st.session_state["logged_in"]):
        if not st.session_state["logged_in"]:
            auth_mode = st.radio("પસંદ કરો:", ["લોગિન (Login)", "નવું એકાઉન્ટ"])
            user_input = st.text_input("યુઝરનામ")
            pass_input = st.text_input("પાસવર્ડ", type="password")
            
            if auth_mode == "લોગિન (Login)":
                if st.button("🔓 લોગિન કરો", use_container_width=True, type="primary"):
                    role = verify_user(user_input, pass_input)
                    if role:
                        st.session_state["logged_in"] = True
                        st.session_state["username"] = user_input
                        st.session_state["role"] = role
                        st.rerun()
                    else:
                        st.error("ખોટો યુઝરનામ અથવા પાસવર્ડ!")
            else:
                if st.button("📝 એકાઉન્ટ બનાવો", use_container_width=True):
                    if user_input and pass_input:
                        if register_user(user_input, pass_input):
                            st.success("એકાઉન્ટ બની ગયું! હવે લોગિન કરો.")
                        else:
                            st.error("આ યુઝરનામ ઉપલબ્ધ નથી!")
        else:
            st.success(f"👤 **{st.session_state['username']}**")
            if st.button("🚪 લોગઆઉટ", use_container_width=True):
                st.session_state["logged_in"] = False
                st.session_state["username"] = ""
                st.rerun()

    with st.expander("🌐 મહત્વપૂર્ણ સરકારી લિંક્સ", expanded=True):
        st.link_button("🌐 લીગલ પોર્ટલ (Blogger)", "https://ramavat12.blogspot.com", use_container_width=True)
        st.link_button("🌐 AnyRoR (૭/૧૨, ૮-અ)", "https://anyror.gujarat.gov.in/", use_container_width=True)
        st.link_button("📜 Digital Gujarat પોર્ટલ", "https://www.digitalgujarat.gov.in", use_container_width=True)
        st.link_button("🏛️ iORA Gujarat", "https://iora.gujarat.gov.in/", use_container_width=True)

# ── HERO DASHBOARD ──────────────────────────────────────────────────
st.markdown('<div class="main-title">ઓમ જનસેવા & ઓનલાઈન સોલ્યુશન સેન્ટર</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">તમામ પ્રકારના સરકારી, મહેસૂલી અને ઓનલાઈન કામકાજ માટેનું વન-સ્ટોપ પોર્ટલ</div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)
with col1:
    st.markdown('<div class="card-box"><div class="card-icon">💳</div><div class="card-title">કાર્ડ & આઈડી સેવાઓ</div><div class="card-desc">આધાર, પાનકાર્ડ, આયુષ્માન કાર્ડ અને આભા કાર્ડ સુધારા.</div></div>', unsafe_allow_html=True)
with col2:
    st.markdown('<div class="card-box"><div class="card-icon">📜</div><div class="card-title">દાખલા & રેશનકાર્ડ</div><div class="card-desc">જન્મ-મરણ સુધારા, LC સુધારા અને રેશનકાર્ડ કામકાજ.</div></div>', unsafe_allow_html=True)
with col3:
    st.markdown('<div class="card-box"><div class="card-icon">🏦</div><div class="card-title">PF & ઇન્સ્યોરન્સ</div><div class="card-desc">પીએફ વિડ્રોઅલ, વાહન વીમો અને લીગલ ટાઈપિંગ.</div></div>', unsafe_allow_html=True)

# ── 3D QUICK BUTTONS ───────────────────────────────────────────────
st.write("")
st.markdown("### 🚀 ક્વિક પોર્ટલ ડાયરેક્ટ લિંક્સ")
st.markdown("""
    <div class="btn-3d-grid">
        <a href="https://ramavat12.blogspot.com" target="_blank" class="btn-3d btn-orange">📧 સરકારી ઈમેલ ડિરેક્ટરી</a>
        <a href="https://uidai.gov.in/" target="_blank" class="btn-3d btn-blue">💳 આધાર પોર્ટલ (UIDAI)</a>
        <a href="https://eportal.incometax.gov.in/" target="_blank" class="btn-3d btn-green">🆔 PAN કાર્ડ પોર્ટલ</a>
        <a href="https://unifiedportal-mem.epfindia.gov.in/" target="_blank" class="btn-3d btn-purple">🏦 PF મેમ્બર પોર્ટલ</a>
        <a href="https://anyror.gujarat.gov.in/" target="_blank" class="btn-3d btn-red">📜 AnyRoR ૭/૧૨</a>
    </div>
""", unsafe_allow_html=True)

# ── MAIN SCREEN SEARCHABLE EMAIL SEARCH TOOL ───────────────────────
st.divider()
st.markdown("### 📧 ગુજરાત સરકારી જિલ્લા/તાલુકા ઈમેલ શોધકો (Search Box)")

col_dist, col_off = st.columns(2)

with col_dist:
    selected_district = st.selectbox("૧. જિલ્લો પસંદ કરો (અથવા નામ ટાઈપ કરો):", list(GOVT_DISTRICT_EMAILS.keys()))

with col_off:
    if selected_district:
        office_options = list(GOVT_DISTRICT_EMAILS[selected_district].keys())
        selected_office = st.selectbox("૨. કચેરી / તાલુકો પસંદ કરો:", office_options)

if selected_district and selected_office:
    email_address = GOVT_DISTRICT_EMAILS[selected_district][selected_office]
    st.markdown(f"""
        <div class="email-result-card">
            <h4>📍 {selected_office} ({selected_district})</h4>
            <p><b>સત્તાવાર ઈમેલ આઈડી:</b></p>
            <h3 style="color: #40a9ff; margin: 5px 0;">{email_address}</h3>
            <a href="mailto:{email_address}" style="color: #52c41a; font-weight: bold; text-decoration: underline;">✉️ સીધો ઈમેલ મોકલવા અહીં ક્લિક કરો</a>
        </div>
    """, unsafe_allow_html=True)

st.divider()

# ── SERVICES & TEMPLATES ───────────────────────────────────────────
TEMPLATE_CATEGORIES = {
    "💳 કાર્ડ & આઈડી સેવાઓ": {
        "પાનકાર્ડ સુધારા / નવું પાનકાર્ડ અરજી":
"""અરજી: પાનકાર્ડ (PAN Card) માં સુધારો કરવા / નવું કાર્ડ મેળવવા બાબત.

સાહેબશ્રી,
હું નીચે સહી કરનાર જણાવું છું કે મારી પાસે હયાત પાનકાર્ડમાં નામ / જન્મતારીખ / પિતાનું નામ ખોટું હોય, તેને આધારકાર્ડ મુજબ સુધારી આપવા વિનંતી છે.

આધારકાર્ડ નકલ તથા જૂના પાનકાર્ડની નકલ સાથે જોડેલ છે.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",

        "આયુષ્માન કાર્ડ / આભા કાર્ડ અરજી":
"""અરજી: પ્રધાનમંત્રી જન આરોગ્ય યોજના (આયુષ્માન ભારત) કાર્ડ મેળવવા બાબત.

સાહેબશ્રી,
અમારી પાસે પાત્રતા ધરાવતું રેશનકાર્ડ / PM ચિઠ્ઠી ઉપલબ્ધ છે. કુટુંબના તમામ સભ્યોના આધારકાર્ડ તથા ઈ-કેવાયસી (e-KYC) ના આધારે આયુષ્માન કાર્ડ જનરેટ કરી આપવા નમ્ર વિનંતી.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",
    },

    "📜 દાખલા & રેશનકાર્ડ સેવાઓ": {
        "જન્મ / મરણ પ્રમાણપત્ર સુધારા અરજી":
"""પ્રતિ,
શ્રીમાન રજિસ્ટ્રાર સાહેબશ્રી (જન્મ-મરણ વિભાગ),
મહાનગરપાલિકા / ગ્રામ પંચાયત: ____________________

વિષય: જન્મ / મરણ રજિસ્ટરમાં નામ / જન્મતારીખ સુધારો કરવા બાબત.
નોંધણી નંબર: ____________, તારીખ: __/__/____

સાહેબશ્રી,
મારા/મારા બાળકના જન્મ/મરણ રજિસ્ટરમાં ભૂલથી નામ કે તારીખ ખોટી લખાયેલ છે:
૧. હાલની ખોટી વિગત: ____________________________________
૨. સાચી વિગત (જે સુધારવાની છે): ____________________________________

સાથે સોગંદનામું, એલસી અને આધારકાર્ડ જોડેલ છે. સુધારીને નવું પ્રમાણપત્ર આપવા વિનંતી.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",

        "રેશનકાર્ડમાં નામ ઉમેરવા / કમી કરવા અરજી":
"""પ્રતિ,
શ્રીમાન પુરવઠા અધિકારી સાહેબશ્રી / મામલતદાર સાહેબશ્રી,
તાલુકા કચેરી: ____________________

વિષય: રેશનકાર્ડ નંબર: ____________ માં નામ ઉમેરવા / કમી કરવા બાબત.

સાહેબશ્રી,
અમારા રેશનકાર્ડમાં નવા સભ્યનું નામ ઉમેરવા / લગ્ન-અવસાનના કારણે નામ કમી કરવા માટેની જરૂરી વિગતો રજૂ કરેલ છે. 

સાથે આધારકાર્ડ તથા આધાર પુરાવા સામેલ છે. ઘટતી કાર્યવાહી કરવા વિનંતી.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",
    },

    "🏦 PF & ઇન્સ્યોરન્સ સેવાઓ": {
        "EPF પીએફ ઉપાડ / કલેમ અરજી":
"""અરજી: એમ્પ્લોઈઝ પ્રોવિડન્ટ ફંડ (EPF) માંથી એડવાન્સ / ફાઈનલ પીએફ ઉપાડ બાબત.

સાહેબશ્રી,
મારો UAN નંબર: ____________________ છે. મારે ઘર કામ/બીમારી/તબીબી સારવાર અર્થે પીએફ માંથી એડવાન્સ રકમ ઉપાડવી હોય, ઓનલાઈન કલેમ સબમિટ કરવા વિનંતી.

સાથે બેંક પાસબુક અને કેન્સલ ચેક જોડેલ છે.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",
    },

    "🌾 મહેસૂલી & ૭/૧૨ સેવાઓ": {
        "૭/૧૨ અને ૮-અ માં વારસાઈ નોંધ દાખલ અરજી":
"""પ્રતિ,
શ્રીમાન મામલતદાર સાહેબશ્રી,
તાલુકા સેવા સદન: ____________________

વિષય: મોજે ગામ: ________, સર્વે/બ્લોક નંબર: ________ માં વારસાઈ નોંધ દાખલ કરવા બાબત.

સાહેબશ્રી,
ઉપરોક્ત જમીનના મૂળ ખાતેદાર શ્રી ____________________________________ નું અવસાન થયેલ હોય, તેઓના કાયદેસરના વારસદારોના નામ ૭/૧૨ અને ૮-અ માં ચડાવવા વિનંતી છે.

સાથે પેઢીનામું, મરણ દાખલો અને વારસદારોના આધારકાર્ડ રજૂ કરેલ છે.

સ્થળ: ____________
તારીખ: __/__/૨૦૨૬""",
    }
}

TEMPLATES = {}
for _cat, _docs in TEMPLATE_CATEGORIES.items():
    TEMPLATES.update(_docs)

# ── WORD GENERATOR ─────────────────────────────────────────────────
def generate_docx(office_name, selected_doc, service_category, reference_no, applicant_name, contact_no, doc_body):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_head = p_head.add_run(f"અરજી: {selected_doc}\n")
    run_head.bold = True
    run_head.font.size = Pt(14)

    p_party = doc.add_paragraph()
    p_party.paragraph_format.line_spacing = 1.5
    run_party = p_party.add_run(f"કચેરી/વિભાગ: {office_name}\nઅરજદારનું નામ: {applicant_name}\nસંપર્ક નંબર: {contact_no}\nસંદર્ભ/આધાર નંબર: {reference_no}\n")
    run_party.font.size = Pt(12)

    p_body = doc.add_paragraph()
    p_body.paragraph_format.line_spacing = 1.5
    run_body = p_body.add_run(doc_body)
    run_body.font.size = Pt(13)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.paragraph_format.space_before = Pt(30)
    run_sign = p_sign.add_run(f"\n\n_____________________\n({applicant_name})\nઅરજદારની સહી")
    run_sign.font.size = Pt(13)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

# ── TABS ───────────────────────────────────────────────────────────
tab1, tab2 = st.tabs(["📝 નવી અરજી બનાવો", "📁 સાચવેલા દસ્તાવેજો"])

with tab1:
    st.subheader("📝 ૧. સેવા શ્રેણી અને અરજી પસંદ કરો")

    if "doc_category" not in st.session_state:
        st.session_state.doc_category = list(TEMPLATE_CATEGORIES.keys())[0]

    cat_cols = st.columns(len(TEMPLATE_CATEGORIES))
    for col, cat_name in zip(cat_cols, TEMPLATE_CATEGORIES.keys()):
        with col:
            is_sel = st.session_state.doc_category == cat_name
            if st.button(cat_name, key=f"cat_{cat_name}", use_container_width=True, type="primary" if is_sel else "secondary"):
                st.session_state.doc_category = cat_name
                st.rerun()

    current_cat_docs = TEMPLATE_CATEGORIES[st.session_state.doc_category]
    doc_names = list(current_cat_docs.keys())

    if "selected_doc" not in st.session_state or st.session_state.selected_doc not in TEMPLATES:
        st.session_state.selected_doc = doc_names[0]
    if st.session_state.selected_doc not in doc_names:
        st.session_state.selected_doc = doc_names[0]

    st.write("")
    doc_cols = st.columns(2)
    for i, doc_name in enumerate(doc_names):
        with doc_cols[i % 2]:
            is_sel_doc = st.session_state.selected_doc == doc_name
            if st.button(doc_name, key=f"doc_{doc_name}", use_container_width=True, type="primary" if is_sel_doc else "secondary"):
                st.session_state.selected_doc = doc_name
                st.rerun()

    selected_doc = st.session_state.selected_doc
    st.success(f"✅ પસંદ કરેલ અરજી: **{selected_doc}**")

    st.divider()

    col_input, col_preview = st.columns([1.1, 0.9], gap="large")

    with col_input:
        st.subheader("🖊️ ગ્રાહક અને કચેરીની વિગત")

        office_name = st.text_input("કચેરી / સબ્જેક્ટ વિગત:", "જનસેવા કેન્દ્ર / મામલતદાર કચેરી / સંબંધિત વિભાગ")

        c1, c2 = st.columns(2)
        with c1:
            applicant_name = st.text_input("અરજદારનું નામ:", "અરજદારનું નામ")
        with c2:
            contact_no = st.text_input("મોબાઈલ નંબર:", "૯૮૭૬૫XXXXX")

        reference_no = st.text_input("આધાર / રેશનકાર્ડ / સંદર્ભ નંબર:", "આધાર / પાન / રેશનકાર્ડ નંબર")

        st.subheader("📄 અરજીનું લખાણ")
        doc_body = st.text_area("મુખ્ય લખાણ:", value=TEMPLATES[selected_doc], height=300, key=f"body_{selected_doc}")

        if st.session_state["logged_in"]:
            if st.button("💾 એકાઉન્ટમાં સેવ કરો", use_container_width=True, type="primary"):
                save_user_draft(
                    st.session_state["username"], selected_doc, st.session_state.doc_category,
                    office_name, applicant_name, contact_no, reference_no, doc_body
                )
                st.success("✅ અરજી સેવ થઈ ગઈ!")
        else:
            st.info("ℹ️ અરજી સેવ કરવા માટે સાઇડબારમાંથી લોગિન કરો.")

    with col_preview:
        st.subheader("👁️ પ્રિવ્યૂ અને ડાઉનલોડ")

        docx_data = generate_docx(office_name, selected_doc, st.session_state.doc_category, reference_no, applicant_name, contact_no, doc_body)

        st.download_button(
            label="📥 ૧. Word (.docx) ફાઇલ ડાઉનલોડ કરો",
            data=docx_data,
            file_name=f"{selected_doc.split(' ')[0]}_Application.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

        st.divider()
        st.subheader("📄 ૨. ડાયરેક્ટ PDF પ્રિન્ટ કરો")

        formatted_body = doc_body.replace('\n', '<br/>')
        print_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                .btn {{ background-color: #1890ff; color: white; padding: 12px; border: none; border-radius: 6px; font-size: 16px; cursor: pointer; width: 100%; font-weight: bold; }}
            </style>
        </head>
        <body>
            <button class="btn" onclick="printDoc()">🖨️ PDF પ્રિન્ટ / સેવ કરો</button>
            <script>
                function printDoc() {{
                    var printWindow = window.open('', '', 'height=600,width=800');
                    printWindow.document.write('<html><head><title>Print Application</title>');
                    printWindow.document.write('<style>body {{ font-family: sans-serif; font-size: 15px; line-height: 1.8; color: #000; padding: 30px; }} .right {{ text-align: right; margin-top: 40px; }}</style></head><body>');
                    printWindow.document.write('<div><b>કચેરી/વિભાગ:</b> {office_name}</div><br/>');
                    printWindow.document.write('<div><b>અરજદાર:</b> {applicant_name} | <b>મોબાઈલ:</b> {contact_no}</div>');
                    printWindow.document.write('<div><b>સંદર્ભ નંબર:</b> {reference_no}</div><hr/>');
                    printWindow.document.write('<div>{formatted_body}</div>');
                    printWindow.document.write('<div class="right">_____________________<br/>({applicant_name})<br/><b>અરજદારની સહી</b></div>');
                    printWindow.document.write('</body></html>');
                    printWindow.document.close();
                    printWindow.focus();
                    setTimeout(function() {{ printWindow.print(); }}, 500);
                }}
            </script>
        </body>
        </html>
        """
        components.html(print_html, height=80)

with tab2:
    st.subheader("🔒 તમારા સેવ કરેલા ડ્રાફ્ટ")
    if not st.session_state["logged_in"]:
        st.warning("🔒 સાચવેલા દસ્તાવેજો જોવા માટે સાઇડબારમાંથી **લોગિન** કરો.")
    else:
        is_admin = (st.session_state["role"] == "admin")
        search_q = st.text_input("🔍 શોધો:", "")
        drafts = get_user_drafts(st.session_state["username"], search_query=search_q, is_admin=is_admin)

        if not drafts:
            st.info("કોઈ સાચવેલા ડ્રાફ્ટ મળ્યા નથી.")
        else:
            for d in drafts:
                draft_id, user_owner, d_type, s_cat, off_name, app_name, cont_no, ref_no, d_body, c_time = d
                with st.expander(f"📜 {d_type} | અરજદાર: {app_name} | ({c_time})"):
                    st.write(f"**વિભાગ:** {s_cat} | **યુઝર:** {user_owner}")
                    st.write(f"**કચેરી:** {off_name} | **મોબાઈલ:** {cont_no} | **સંદર્ભ:** {ref_no}")
                    st.text_area("લખાણ:", value=d_body, height=150, key=f"text_{draft_id}")

                    c_dl, c_del = st.columns(2)
                    with c_dl:
                        saved_docx = generate_docx(off_name, d_type, s_cat, ref_no, app_name, cont_no, d_body)
                        st.download_button("📥 Word ડાઉનલોડ", data=saved_docx, file_name=f"Draft_{draft_id}.docx", key=f"dl_{draft_id}")
                    with c_del:
                        if st.button("🗑️ ડિલીટ કરો", key=f"del_{draft_id}"):
                            delete_user_draft(draft_id, st.session_state["username"], is_admin=is_admin)
                            st.success("દસ્તાવેજ ડિલીટ થઈ ગયો!")
                            st.rerun()
